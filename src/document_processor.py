"""
Document Processor Module

This module handles loading English documents from .docx files,
splitting them into chunks, and creating a vector store for RAG.
"""

import os
import sys
from pathlib import Path
from typing import List
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_core.documents import Document as LangChainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from dotenv import load_dotenv

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent))
from contact_loader import load_contacts_as_document

# Load environment variables
load_dotenv()


class DocumentProcessor:
    """Process documents for RAG-based chatbot"""
    
    def __init__(
        self,
        docs_dir: str = "data/docs/national-bank-products",
        vectorstore_dir: str = "data/vectorstore",
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        """
        Initialize the document processor
        
        Args:
            docs_dir: Directory containing .docx files
            vectorstore_dir: Directory to store vector database
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.docs_dir = Path(docs_dir)
        self.vectorstore_dir = Path(vectorstore_dir)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Initialize OpenAI embeddings
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small"
        )
        
        self.vectorstore = None
    
    def load_docx(self, file_path: Path) -> str:
        """
        Extract text from a .docx file
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Extracted text content
        """
        doc = DocxDocument(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        return "\n\n".join(full_text)
    
    def load_pdf(self, file_path: Path) -> str:
        """
        Extract text from a PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        reader = PdfReader(file_path)
        full_text = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                full_text.append(text)
        
        return "\n\n".join(full_text)
    
    def load_all_documents(self) -> List[LangChainDocument]:
        """
        Load all .docx and .pdf files from the documents directory,
        plus contact information
        
        Returns:
            List of LangChain Document objects
        """
        documents = []
        
        if not self.docs_dir.exists():
            raise FileNotFoundError(f"Documents directory not found: {self.docs_dir}")
        
        # Get both PDF and DOCX files
        pdf_files = list(self.docs_dir.glob("*.pdf"))
        docx_files = list(self.docs_dir.glob("*.docx"))
        all_files = pdf_files + docx_files
        
        if not all_files:
            raise FileNotFoundError(f"No .pdf or .docx files found in {self.docs_dir}")
        
        print(f"Found {len(all_files)} document(s) to process ({len(pdf_files)} PDFs, {len(docx_files)} DOCX)...")
        
        for file_path in all_files:
            try:
                print(f"Loading: {file_path.name}")
                
                # Choose loader based on file extension
                if file_path.suffix.lower() == '.pdf':
                    text_content = self.load_pdf(file_path)
                else:
                    text_content = self.load_docx(file_path)
                
                # Create LangChain Document with metadata
                doc = LangChainDocument(
                    page_content=text_content,
                    metadata={
                        "source": file_path.name,
                        "file_path": str(file_path),
                        "file_type": file_path.suffix.lower()
                    }
                )
                documents.append(doc)
                
            except Exception as e:
                print(f"Error loading {file_path.name}: {e}")
                continue
        
        # Add contact information
        print("Loading contact information...")
        contact_docs = load_contacts_as_document()
        documents.extend(contact_docs)
        print(f"Added {len(contact_docs)} contact document(s)")
        
        print(f"Successfully loaded {len(documents)} document(s) total")
        return documents
    
    def split_documents(self, documents: List[LangChainDocument]) -> List[LangChainDocument]:
        """
        Split documents into smaller chunks
        
        Args:
            documents: List of documents to split
            
        Returns:
            List of document chunks
        """
        print(f"Splitting documents into chunks (size={self.chunk_size}, overlap={self.chunk_overlap})...")
        chunks = self.text_splitter.split_documents(documents)
        print(f"Created {len(chunks)} chunks")
        return chunks
    
    def create_vectorstore(self, chunks: List[LangChainDocument]) -> Chroma:
        """
        Create a vector store from document chunks
        
        Args:
            chunks: List of document chunks
            
        Returns:
            Chroma vector store
        """
        print("Creating vector store with OpenAI embeddings...")
        
        # Create vectorstore directory if it doesn't exist
        self.vectorstore_dir.mkdir(parents=True, exist_ok=True)
        
        # Create Chroma vector store
        vectorstore = Chroma.from_documents(
            documents=chunks,
            embedding=self.embeddings,
            persist_directory=str(self.vectorstore_dir)
        )
        
        print(f"Vector store created and saved to {self.vectorstore_dir}")
        return vectorstore
    
    def load_vectorstore(self) -> Chroma:
        """
        Load an existing vector store
        
        Returns:
            Chroma vector store
        """
        if not self.vectorstore_dir.exists():
            raise FileNotFoundError(f"Vector store not found at {self.vectorstore_dir}")
        
        print(f"Loading vector store from {self.vectorstore_dir}...")
        vectorstore = Chroma(
            persist_directory=str(self.vectorstore_dir),
            embedding_function=self.embeddings
        )
        
        return vectorstore
    
    def process_documents(self, force_recreate: bool = False):
        """
        Complete pipeline: load documents, split, and create vector store
        
        Args:
            force_recreate: If True, recreate vector store even if it exists
        """
        # Check if vector store already exists
        if self.vectorstore_dir.exists() and not force_recreate:
            print("Vector store already exists. Loading existing store...")
            self.vectorstore = self.load_vectorstore()
            return self.vectorstore
        
        # Load documents
        documents = self.load_all_documents()
        
        # Split into chunks
        chunks = self.split_documents(documents)
        
        # Create vector store
        self.vectorstore = self.create_vectorstore(chunks)
        
        return self.vectorstore


if __name__ == "__main__":
    # Test the document processor
    processor = DocumentProcessor()
    
    try:
        # Process documents and create vector store
        vectorstore = processor.process_documents(force_recreate=False)
        
        # Test retrieval
        print("\n" + "="*50)
        print("Testing retrieval with sample query...")
        print("="*50)
        
        test_query = "What happened in the news?"
        results = vectorstore.similarity_search(test_query, k=2)
        
        print(f"\nQuery: '{test_query}'")
        print(f"\nTop {len(results)} results:")
        for i, doc in enumerate(results, 1):
            print(f"\n--- Result {i} ---")
            print(f"Source: {doc.metadata['source']}")
            print(f"Content preview: {doc.page_content[:200]}...")
        
    except Exception as e:
        print(f"Error: {e}")
